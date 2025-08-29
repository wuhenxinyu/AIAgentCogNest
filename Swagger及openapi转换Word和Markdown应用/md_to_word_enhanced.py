"""
Markdown转Word增强工具：先将Markdown转换为HTML，再将HTML转换为Word文档
支持更完整的样式保留和格式转换
"""
import os
import argparse
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def convert_markdown_to_html(md_content):
    """将Markdown内容转换为HTML"""
    # 使用markdown库将Markdown转换为HTML
    html_content = markdown.markdown(md_content, extensions=[
        'extra',  # 额外功能，如表格、代码块等
        'toc',    # 目录生成
        'sane_lists',  # 更合理的列表处理
        'smarty'  # 智能引号等
    ])
    return html_content


def convert_html_to_word(html_content, output_docx_path):
    """将HTML内容转换为Word文档"""
    # 创建Word文档
    doc = Document()
    
    # 解析HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 处理文档标题（如果有）
    title_tag = soup.find('h1')
    if title_tag:
        title_text = title_tag.get_text().strip()
        # 设置文档标题
        title_heading = doc.add_heading(title_text, level=0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 移除已处理的标题
        title_tag.decompose()
    
    # 当前段落样式跟踪
    current_style = None
    
    # 遍历HTML内容的所有顶级元素
    for element in soup.find_all(recursive=False):
        # 处理标题
        if element.name.startswith('h') and len(element.name) == 2:
            level = int(element.name[1])
            text = element.get_text().strip()
            
            # 根据级别添加相应的标题
            if 1 <= level <= 6:
                # Word的heading级别是0-4，需要映射
                word_level = min(4, level - 1)
                heading = doc.add_heading(text, level=word_level)
                
                # 尝试保留原始标题的样式（如颜色、大小等）
                if element.get('style'):
                    # 这里简化处理，实际可能需要更复杂的样式解析
                    pass
        
        # 处理段落
        elif element.name == 'p':
            text = element.get_text().strip()
            if text:
                paragraph = doc.add_paragraph(text)
                
                # 检查段落是否有特殊样式
                if element.get('align') == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif element.get('align') == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 处理列表（有序和无序）
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                text = li.get_text().strip()
                if text:
                    # 根据列表类型选择样式
                    style = 'List Bullet' if element.name == 'ul' else 'List Number'
                    doc.add_paragraph(text, style=style)
        
        # 处理代码块
        elif element.name == 'pre':
            code_block = element.find('code')
            if code_block:
                code_text = code_block.get_text()
                
                # 添加代码块标题（尝试获取语言信息）
                language = ''
                if code_block.get('class'):
                    for cls in code_block.get('class'):
                        if cls.startswith('language-'):
                            language = cls[9:]
                            break
                
                if language:
                    doc.add_paragraph(f"代码语言: {language}")
                
                # 添加代码内容，使用等宽字体
                code_paragraph = doc.add_paragraph()
                code_run = code_paragraph.add_run(code_text)
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(10)
        
        # 处理表格
        elif element.name == 'table':
            # 获取表头
            thead = element.find('thead')
            tbody = element.find('tbody')
            
            if thead and tbody:
                # 获取表头行
                header_row = thead.find('tr')
                if header_row:
                    # 获取列数
                    header_cells = header_row.find_all(['th', 'td'])
                    col_count = len(header_cells)
                    
                    # 创建表格
                    table = doc.add_table(rows=1, cols=col_count)
                    
                    # 添加表头
                    for i, cell in enumerate(header_cells):
                        table.rows[0].cells[i].text = cell.get_text().strip()
                        # 设置表头样式（加粗）
                        for paragraph in table.rows[0].cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # 添加表格数据
                    for row in tbody.find_all('tr'):
                        data_cells = row.find_all(['th', 'td'])
                        if data_cells:
                            new_row = table.add_row()
                            for i, cell in enumerate(data_cells[:col_count]):  # 防止列数不匹配
                                new_row.cells[i].text = cell.get_text().strip()
        
        # 处理引用块
        elif element.name == 'blockquote':
            text = element.get_text().strip()
            if text:
                paragraph = doc.add_paragraph(text)
                # 设置引用样式（这里简化处理）
                paragraph.paragraph_format.left_indent = Inches(0.5)
        
        # 处理水平线
        elif element.name == 'hr':
            doc.add_page_break()  # 用分页符代替水平线
        
        # 处理其他内容
        else:
            # 对于未专门处理的元素，尝试提取文本内容
            text = element.get_text().strip()
            if text:
                doc.add_paragraph(text)
    
    # 保存Word文档
    doc.save(output_docx_path)
    return True


def convert_md_to_word_through_html(md_file_path, output_docx_path):
    """通过HTML中间步骤将Markdown文件转换为Word文档"""
    # 检查输入文件是否存在
    if not os.path.exists(md_file_path):
        print(f"错误：Markdown文件 '{md_file_path}' 不存在")
        return False
    
    try:
        print(f"正在读取Markdown文件：{md_file_path}")
        # 读取Markdown文件内容
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        print("正在将Markdown转换为HTML...")
        # 将Markdown转换为HTML
        html_content = convert_markdown_to_html(md_content)
        
        print("正在将HTML转换为Word文档...")
        # 将HTML转换为Word文档
        success = convert_html_to_word(html_content, output_docx_path)
        
        if success:
            print(f"Word文档已成功保存到：{output_docx_path}")
            return True
        else:
            print("HTML转Word失败")
            return False
            
    except Exception as e:
        print(f"转换过程中发生错误：{str(e)}")
        return False


def main():
    # 命令行参数解析
    parser = argparse.ArgumentParser(description='通过HTML中间步骤将Markdown文件转换为Word文档')
    parser.add_argument('-i', '--input', required=True, help='输入的Markdown文件路径')
    parser.add_argument('-o', '--output', help='输出的Word文件路径（默认与输入文件同名，后缀为.docx）')
    
    args = parser.parse_args()
    
    # 确定输出文件路径
    if not args.output:
        base_name = os.path.splitext(os.path.basename(args.input))[0]
        output_path = os.path.join(os.path.dirname(args.input or '.'), f"{base_name}.docx")
    else:
        output_path = args.output
    
    # 执行转换
    convert_md_to_word_through_html(args.input, output_path)


if __name__ == "__main__":
    main()