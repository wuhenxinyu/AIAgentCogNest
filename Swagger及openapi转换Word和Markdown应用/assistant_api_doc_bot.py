"""
API文档编写助手
  1、使用Qwen-agent封装swagger_to_markdown和md_to_word_enhanced功能
  2、支持将Swagger/OpenAPI JSON转换为Markdown，再转换为Word文档
  3、提供Web界面和终端界面两种交互方式
"""
import os
import asyncio
from typing import Optional
import dashscope
from qwen_agent.agents import Assistant
from qwen_agent.gui import WebUI
from qwen_agent.tools.base import BaseTool, register_tool
import json
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import time

# 配置 DashScope
dashscope.api_key = os.getenv('DASHSCOPE_API_KEY', 'sk-58f051ae745e4bb19fdca31735105b11')  # 从环境变量获取 API Key
dashscope.timeout = 30  # 设置超时时间为 30 秒
# os.environ['QWEN_AGENT_DEFAULT_WORKSPACE'] = 'workspace'
os.environ['QWEN_AGENT_DEFAULT_MAX_INPUT_TOKENS'] = '158000'
# os.environ['GRADIO_TEMP_DIR'] = '/Users/clz/Downloads/xmkf/llm/AIAgentCogNest/upload/'

# ====== API文档编写助手 system prompt 和函数描述 ======
system_prompt = """我是API文档编写助手，我可以帮助您将Swagger/OpenAPI JSON文件转换为格式良好的Word文档。

我提供以下功能：
1. 将Swagger/OpenAPI JSON文件转换为Markdown格式的API文档
2. 将Markdown文档转换为格式化的Word文档（.docx）
3. 支持完整的样式保留和格式转换
4. 处理标题、表格、代码块等多种元素

您可以直接提供Swagger JSON文件路径，我会使用工具为您进行转换。

请注意：目前暂不支持PDF和TXT格式文件的直接转换。如果您需要处理这些格式的内容，建议您先将其转换为Markdown格式，然后再使用我提供的Markdown转Word文档功能进行处理。
"""

functions_desc = [
    {
        "name": "swagger_to_docx",
        "description": "将Swagger/OpenAPI JSON文件转换为Word文档",
        "parameters": {
            "type": "object",
            "properties": {
                "swagger_file_path": {
                    "type": "string",
                    "description": "Swagger/OpenAPI JSON文件的路径",
                },
                "output_file_path": {
                    "type": "string",
                    "description": "输出Word文档的路径（可选）",
                }
            },
            "required": ["swagger_file_path"],
        },
    },
    {
        "name": "markdown_to_docx",
        "description": "将Markdown文件转换为Word文档",
        "parameters": {
            "type": "object",
            "properties": {
                "md_file_path": {
                    "type": "string",
                    "description": "Markdown文件的路径",
                },
                "output_file_path": {
                    "type": "string",
                    "description": "输出Word文档的路径（可选）",
                }
            },
            "required": ["md_file_path"],
        },
    },
    {
        "name": "swagger_to_markdown",
        "description": "将Swagger/OpenAPI JSON文件转换为Markdown文档",
        "parameters": {
            "type": "object",
            "properties": {
                "swagger_file_path": {
                    "type": "string",
                    "description": "Swagger/OpenAPI JSON文件的路径",
                },
                "output_file_path": {
                    "type": "string",
                    "description": "输出Markdown文档的路径（可选）",
                }
            },
            "required": ["swagger_file_path"],
        },
    },
]

# ====== swagger_to_markdown 工具类实现 ======
@register_tool('swagger_to_markdown')
class SwaggerToMarkdownTool(BaseTool):
    """
    Swagger/OpenAPI JSON转Markdown工具
    """
    description = '将Swagger/OpenAPI JSON文件转换为Markdown文档'
    parameters = [{
        'name': 'swagger_file_path',
        'type': 'string',
        'description': 'Swagger/OpenAPI JSON文件的路径',
        'required': True
    }, {
        'name': 'output_file_path',
        'type': 'string',
        'description': '输出Markdown文档的路径（可选）',
        'required': False
    }]

    def call(self, params: str, **kwargs) -> str:
        try:
            args = json.loads(params)
            swagger_file_path = args['swagger_file_path']
            output_file_path = args.get('output_file_path')
            
            # 检查输入文件是否存在，如果不存在则尝试在GRADIO_TEMP_DIR目录下遍历查找
            if not os.path.exists(swagger_file_path):
                # 获取GRADIO_TEMP_DIR目录
                gradio_temp_dir = os.environ.get('GRADIO_TEMP_DIR', '')
                if gradio_temp_dir and os.path.exists(gradio_temp_dir):
                    # 获取要查找的文件名
                    target_filename = os.path.basename(swagger_file_path)
                    found_paths = []
                    
                    # 遍历GRADIO_TEMP_DIR目录及其所有子目录
                    for root, dirs, files in os.walk(gradio_temp_dir):
                        if target_filename in files:
                            found_paths.append(os.path.join(root, target_filename))
                    
                    # 根据查找结果处理
                    if len(found_paths) == 1:
                        swagger_file_path = found_paths[0]
                    elif len(found_paths) > 1:
                        return f"错误：在GRADIO_TEMP_DIR目录下找到多个同名文件：{', '.join(found_paths)}"
                    else:
                        return f"错误：Swagger文件 '{swagger_file_path}' 不存在，且在GRADIO_TEMP_DIR目录下也未找到"
                else:
                    return f"错误：Swagger文件 '{swagger_file_path}' 不存在"
            
            # 读取Swagger JSON文件
            with open(swagger_file_path, 'r', encoding='utf-8') as f:
                swagger_data = json.load(f)
            
            # 生成输出文件路径
            if not output_file_path:
                base_name = os.path.splitext(os.path.basename(swagger_file_path))[0]
                output_file_path = f"{base_name}_api文档.md"
            
            # 存储Markdown内容
            md_content = []
            
            # 获取数据模型定义（支持OpenAPI 3.0和Swagger 2.0）
            definitions = swagger_data.get('definitions', {})  # Swagger 2.0
            if not definitions:
                definitions = swagger_data.get('components', {}).get('schemas', {})  # OpenAPI 3.0
            
            paths = swagger_data.get('paths', {})
            if paths:
                # 按路径分组
                for path, path_item in paths.items():
                    for method, operation in path_item.items():
                        md_content.append(f" ")
                        md_content.append(f" ")
                        # 方法名称转换为大写
                        method = method.upper()
                        
                        # 获取接口信息
                        summary = operation.get('summary', '')
                        desc = operation.get('description', '')
                        tags = operation.get('tags', [])
                        consumes = operation.get('consumes', [])
                        produces = operation.get('produces', [])
                        
                        # 添加接口标题和锚点
                        interface_title = summary if summary else f"{method} {path}"
                        md_content.append(f"### 名称：{interface_title}")

                        # 1. 在读取Swagger数据后获取basePath
                        # 移除重复读取文件的代码，直接使用已加载的 swagger_data
                        # 获取basePath（如果存在）
                        base_path = swagger_data.get('basePath', '')
                        
                        # 2. 在生成请求路径时拼接basePath
                        md_content.append(f"- **请求路径**: https://xxx.xxx.xxx.xxx{base_path}{path}")
                        md_content.append(f"- **请求方法**: {method}")
                        md_content.append(f"- **接口描述**: {desc if desc else '暂无描述'}")
                        
                        # 添加标签
                        if tags:
                            md_content.append(f"- **接口标签**: {', '.join(tags)}")
                            
                        md_content.append("- **请求参数**:")
                        # 请求参数
                        parameters = operation.get('parameters', [])
                        if parameters:
                            # 按参数位置分类
                            query_params = []
                            header_params = []
                            path_params = []
                            body_params = []
                            # 新增：存储需要展示的请求参数结构体详情
                            request_body_schemas = []
                            
                            for param in parameters:
                                param_location = param.get('in', 'query')
                                if param_location == 'query':
                                    query_params.append(param)
                                elif param_location == 'header':
                                    header_params.append(param)
                                elif param_location == 'path':
                                    path_params.append(param)
                                elif param_location == 'body':
                                    body_params.append(param)
                            
                            # 显示Query参数
                            if query_params:
                                md_content.append("<table>")
                                md_content.append("  <thead class='ant-table-thead'>")
                                md_content.append("    <tr>")
                                md_content.append("      <th key=name>参数名称</th><th key=type>类型</th><th key=required>是否必须</th><th key=default>默认值</th><th key=example>示例</th><th key=desc>备注</th>")
                                md_content.append("    </tr>")
                                # 在第87行附近修改Query参数表格的tbody属性
                                md_content.append("  </thead><tbody class='ant-table-tbody'>")
                                
                                for idx, param in enumerate(query_params):
                                    name = param.get('name', '')
                                    required = '是' if param.get('required', False) else '否'
                                    example = param.get('example', '') if 'example' in param else ''
                                    desc = param.get('description', '')
                                    default_value = param.get('default', '')
                                    
                                    # 获取参数类型和长度信息
                                    param_type = param.get('type', 'string')
                                    if 'format' in param:
                                        param_type = f"{param_type}({param['format']})"
                                    
                                    # 添加长度信息
                                    length_info = []
                                    if 'maxLength' in param:
                                        length_info.append(f"最大长度: {param['maxLength']}")
                                    if 'minLength' in param:
                                        length_info.append(f"最小长度: {param['minLength']}")
                                    if 'maximum' in param:
                                        length_info.append(f"最大值: {param['maximum']}")
                                    if 'minimum' in param:
                                        length_info.append(f"最小值: {param['minimum']}")
                                    
                                    # 合并类型和长度信息
                                    full_type_info = param_type
                                    if length_info:
                                        full_type_info += " (" + ", ".join(length_info) + ")"
                                    
                                    md_content.append(f"    <tr key={idx}>")
                                    md_content.append(f"      <td key=0>{name}</td>")
                                    md_content.append(f"      <td key=1>{full_type_info}</td>")
                                    md_content.append(f"      <td key=2>{required}</td>")
                                    md_content.append(f"      <td key=3>{default_value}</td>")
                                    md_content.append(f"      <td key=4>{example}</td>")
                                    md_content.append(f"      <td key=5>{desc}</td>")
                                    md_content.append(f"    </tr>")
                                
                                md_content.append("  </tbody>")
                                md_content.append("</table>")
                            
                            # 显示Header参数
                            if header_params:
                                md_content.append("<table>")
                                md_content.append("  <thead class='ant-table-thead'>")
                                md_content.append("    <tr>")
                                md_content.append("      <th key=name>参数名称</th><th key=type>类型</th><th key=required>是否必须</th><th key=default>默认值</th><th key=example>示例</th><th key=desc>备注</th>")
                                md_content.append("    </tr>")
                                # 修正 className 为 class
                                md_content.append("  </thead><tbody class='ant-table-tbody'>")
                                
                                for idx, param in enumerate(header_params):
                                    name = param.get('name', '')
                                    required = '是' if param.get('required', False) else '否'
                                    desc = param.get('description', '')
                                    default_value = param.get('default', '')
                                    # 修复：添加 example 变量定义
                                    example = param.get('example', '') if 'example' in param else ''
                                    
                                    # 获取参数类型和长度信息
                                    param_type = param.get('type', 'string')
                                    if 'format' in param:
                                        param_type = f"{param_type}({param['format']})"
                                    
                                    # 添加长度信息
                                    length_info = []
                                    if 'maxLength' in param:
                                        length_info.append(f"最大长度: {param['maxLength']}")
                                    if 'minLength' in param:
                                        length_info.append(f"最小长度: {param['minLength']}")
                                    
                                    # 合并类型和长度信息
                                    full_type_info = param_type
                                    if length_info:
                                        full_type_info += " (" + ", ".join(length_info) + ")"
                                    
                                    md_content.append(f"    <tr key={idx}>")
                                    md_content.append(f"      <td key=0>{name}</td>")
                                    md_content.append(f"      <td key=1>{full_type_info}</td>")
                                    md_content.append(f"      <td key=2>{required}</td>")
                                    md_content.append(f"      <td key=3>{default_value}</td>")
                                    md_content.append(f"      <td key=4>{example}</td>")
                                    md_content.append(f"      <td key=5>{desc}</td>")
                                    md_content.append(f"    </tr>")
                                
                                md_content.append("  </tbody>")
                                md_content.append("</table>")
                            
                            # 显示Path参数
                            if path_params:
                                md_content.append("<table>")
                                md_content.append("  <thead class='ant-table-thead'>")
                                md_content.append("    <tr>")
                                md_content.append("      <th key=name>参数名称</th><th key=type>类型</th><th key=required>是否必须</th><th key=default>默认值</th><th key=example>示例</th><th key=desc>备注</th>")
                                md_content.append("    </tr>")
                                # 修正 className 为 class
                                md_content.append("  </thead><tbody class='ant-table-tbody'>")
                                
                                for idx, param in enumerate(path_params):
                                    name = param.get('name', '')
                                    required = '是' if param.get('required', False) else '否'
                                    example = param.get('example', '') if 'example' in param else ''
                                    desc = param.get('description', '')
                                    default_value = param.get('default', '')
                                    
                                    # 获取参数类型和长度信息
                                    param_type = param.get('type', 'string')
                                    if 'format' in param:
                                        param_type = f"{param_type}({param['format']})"
                                    
                                    # 添加长度信息
                                    length_info = []
                                    if 'maxLength' in param:
                                        length_info.append(f"最大长度: {param['maxLength']}")
                                    if 'minLength' in param:
                                        length_info.append(f"最小长度: {param['minLength']}")
                                    if 'maximum' in param:
                                        length_info.append(f"最大值: {param['maximum']}")
                                    if 'minimum' in param:
                                        length_info.append(f"最小值: {param['minimum']}")
                                    
                                    # 合并类型和长度信息
                                    full_type_info = param_type
                                    if length_info:
                                        full_type_info += " (" + ", ".join(length_info) + ")"
                                    
                                    md_content.append(f"    <tr key={idx}>")
                                    md_content.append(f"      <td key=0>{name}</td>")
                                    md_content.append(f"      <td key=1>{full_type_info}</td>")
                                    md_content.append(f"      <td key=2>{required}</td>")
                                    md_content.append(f"      <td key=3>{default_value}</td>")
                                    md_content.append(f"      <td key=4>{example}</td>")
                                    md_content.append(f"      <td key=5>{desc}</td>")
                                    md_content.append(f"    </tr>")
                                
                                md_content.append("  </tbody>")
                                md_content.append("</table>")
                            
                            # 显示Body参数
                            if body_params:
                                md_content.append("<table>")
                                md_content.append("  <thead class='ant-table-thead'>")
                                md_content.append("    <tr>")
                                md_content.append("      <th key=name>参数名称</th><th key=type>类型</th><th key=required>是否必须</th><th key=default>默认值</th><th key=example>示例</th><th key=desc>备注</th>")
                                md_content.append("    </tr>")
                                # 在第243行附近修改Body参数表格的tbody属性
                                md_content.append("  </thead><tbody class='ant-table-tbody'>")
                                
                                for idx, param in enumerate(body_params):
                                    name = param.get('name', '')
                                    required = '是' if param.get('required', False) else '否'
                                    desc = param.get('description', '')
                                    default_value = param.get('default', '')
                                    
                                    # 获取参数类型
                                    param_type = 'object'
                                    if 'schema' in param:
                                        schema = param['schema']
                                        # 初始化ref_name变量，防止UnboundLocalError
                                        ref_name = None
                                        
                                        if '$ref' in schema:
                                            ref_name = schema['$ref'].split('/')[-1]
                                            # 添加替换逻辑
                                            ref_name = ref_name.replace('devops_aishu_cn_AISHUDevOps_AnyFabric__git_', '')
                                            # 检查引用的实际类型
                                            if ref_name in definitions:
                                                ref_schema = definitions[ref_name]
                                                actual_type = ref_schema.get('type', 'object')
                                                param_type = f"{ref_name} ({actual_type})"
                                                # 新增：将结构体详情信息存储起来，稍后处理
                                                request_body_schemas.append({"ref_name": ref_name, "schema": ref_schema})
                                            else:
                                                param_type = ref_name
                                    
                                    md_content.append(f"    <tr key={idx}>")
                                    md_content.append(f"      <td key=0>{name}</td>")
                                    md_content.append(f"      <td key=1>{param_type}</td>")
                                    md_content.append(f"      <td key=2>{required}</td>")
                                    md_content.append(f"      <td key=3>{default_value}</td>")
                                    md_content.append(f"      <td key=4>{''}</td>")  # Body 参数示例可能未提供，置空
                                    md_content.append(f"      <td key=5>{desc}</td>")
                                    md_content.append(f"    </tr>")
                                
                                md_content.append("  </tbody>")
                                md_content.append("</table>")
                            
                            # 新增：在所有参数表格处理完成后，单独处理请求参数结构体详情
                            for schema_info in request_body_schemas:
                                ref_name = schema_info["ref_name"]
                                schema = schema_info["schema"]
                                md_content.append(f"\n- **请求参数结构体详情**:\n")
                                md_content.append("<table>")
                                md_content.append("  <thead class='ant-table-thead'>")
                                md_content.append("    <tr>")
                                md_content.append("      <th key=name>字段名称</th><th key=type>类型</th><th key=required>是否必须</th><th key=default>默认值</th><th key=desc>备注</th>")
                                md_content.append("    </tr>")
                                md_content.append("  </thead><tbody class='ant-table-tbody'>")
                                
                                # 使用expand_schema_as_tree函数展开结构体字段
                                expand_schema_as_tree(md_content, schema, definitions, ref_name, "", 0, 0)
                                
                                md_content.append("  </tbody>")
                                md_content.append("</table>")
                                md_content.append("")
                        
                        # 请求体
                        request_body = operation.get('requestBody', {})
                        if request_body:
                            if 'content' in request_body:
                                for content_type, content in request_body['content'].items():
                                    md_content.append(f"**Content-Type**: {content_type}")
                                    md_content.append("<table>")
                                    md_content.append("  <thead class='ant-table-thead'>")
                                    md_content.append("    <tr>")
                                    md_content.append("      <th key=name>名称</th><th key=type>类型</th><th key=required>是否必须</th><th key=default>默认值</th><th key=desc>备注</th>")
                                    md_content.append("    </tr>")
                                    # 修正 className 为 class
                                    md_content.append("  </thead><tbody class='ant-table-tbody'>")
                                    
                                    if 'schema' in content:
                                        schema = content['schema']
                                        if '$ref' in schema:
                                            ref_name = schema['$ref'].split('/')[-1]
                                            if ref_name in definitions:
                                                schema = definitions[ref_name]
                                                # 展开schema结构
                                                expand_schema_as_tree(md_content, schema, definitions, ref_name, "", 0, 0)
                                        else:
                                            # 展开简单schema结构
                                            expand_schema_as_tree(md_content, schema, definitions, "body", "", 0, 0)
                                    
                                    md_content.append("  </tbody>")
                                    md_content.append("</table>")
                        
                        # 响应参数
                        responses = operation.get('responses', {})
                        if responses:
                            md_content.append("- **响应参数**:")
                            
                            # 只选择主要的成功响应来展开（通常是200）
                            primary_response = None
                            for status_code, response in responses.items():
                                if status_code.startswith('2'):  # 成功状态码
                                    primary_response = response
                                    break
                                
                            # 修复缩进：将if条件移到for循环外部
                            if primary_response:
                                md_content.append("<table>")
                                md_content.append("  <thead class='ant-table-thead'>")
                                md_content.append("    <tr>")
                                md_content.append("      <th key=name>名称</th><th key=type>类型</th><th key=required>是否必须</th><th key=default>默认值</th><th key=desc>备注</th>")
                                md_content.append("    </tr>")
                                # 修正 className 为 class
                                md_content.append("  </thead><tbody class='ant-table-tbody'>")
                                
                                if 'content' in primary_response:
                                    for content_type, content in primary_response['content'].items():
                                        if 'schema' in content:
                                            schema = content['schema']
                                            if '$ref' in schema:
                                                ref_name = schema['$ref'].split('/')[-1]
                                                if ref_name in definitions:
                                                    schema = definitions[ref_name]
                                                    # 修复：移除is_response_param参数
                                                    expand_schema_as_tree(md_content, schema, definitions, ref_name, "", 0, 0)
                                            else:
                                                # 修复：移除is_response_param参数
                                                expand_schema_as_tree(md_content, schema, definitions, "response", "", 0, 0)
                                elif 'schema' in primary_response and primary_response['schema']:
                                        # Swagger 2.0格式
                                        # 修复：首先从primary_response中获取schema
                                        schema = primary_response['schema']
                                        
                                        # 检查schema是否包含引用
                                        if '$ref' in schema:
                                            ref_name = schema['$ref'].split('/')[-1]
                                            if ref_name in definitions:
                                                schema = definitions[ref_name]
                                                # 展开schema结构
                                                expand_schema_as_tree(md_content, schema, definitions, ref_name, "", 0, 0)
                                        else:
                                            # 展开简单schema结构
                                            expand_schema_as_tree(md_content, schema, definitions, "response", "", 0, 0)
                                
                                md_content.append("  </tbody>")
                                md_content.append("</table>")
                            
                            # 显示所有状态码的简要说明
                            status_codes_table = False
                            for status_code, response in responses.items():
                                if not status_codes_table:
                                    md_content.append("- **状态码说明**")
                                    md_content.append("<table>")
                                    md_content.append("  <thead class='ant-table-thead'>")
                                    md_content.append("    <tr>")
                                    md_content.append("      <th key=code>状态码</th><th key=desc>描述</th>")
                                    md_content.append("    </tr>")
                                    # 修正 className 为 class
                                    md_content.append("  </thead><tbody class='ant-table-tbody'>")
                                    status_codes_table = True
                                
                                description = response.get('description', '')
                                md_content.append(f"    <tr key={status_code}>")
                                md_content.append(f"      <td key=code>{status_code}</td>")
                                md_content.append(f"      <td key=desc>{description}</td>")
                                md_content.append(f"    </tr>")
                            
                            if status_codes_table:
                                md_content.append("  </tbody>")
                                md_content.append("</table>")
                        
                        md_content.append("")
            
            # 将Markdown内容写入文件
            with open(output_file_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_content))
            
            print(f"Markdown文档已生成: {output_file_path}")
            return f"Swagger文件已成功转换为Markdown文档：{output_file_path}"
        except json.JSONDecodeError:
            return "错误：参数解析失败，请检查输入参数是否为有效的JSON格式"
        except Exception as e:
            return f"Swagger转Markdown失败: {str(e)}"


# 移动expand_schema_as_tree函数到SwaggerToMarkdownTool类外部

def expand_schema_as_tree(md_content, schema, definitions, schema_name=None, parent_path="", indent_level=0, row_index=0):
    """将schema展开为树形结构的HTML表格行"""
    # 处理数组类型
    if schema.get('type') == 'array' and 'items' in schema:
        items_schema = schema['items']
        item_type = 'object'
        
        if '$ref' in items_schema:
            ref_name = items_schema['$ref'].split('/')[-1]
            item_type = ref_name
            if ref_name in definitions:
                items_schema = definitions[ref_name]
        elif 'type' in items_schema:
            item_type = items_schema['type']
            if 'format' in items_schema:
                item_type = f"{item_type}({items_schema['format']})"
        
        # 显示数组行
        padding = indent_level * 20
        md_content.append(f"    <tr key={row_index}-{indent_level}>")
        md_content.append(f'      <td key=0><span style="padding-left: {padding}px"><span style="color: #8c8a8a"></span> {schema_name}</span></td>')
        md_content.append(f"      <td key=1><span>{schema.get('type')} [{item_type}]</span></td>")
        md_content.append(f"      <td key=2>是</td>")
        md_content.append(f"      <td key=3></td>")
        md_content.append(f'      <td key=4><span style="white-space: pre-wrap">{schema.get("description", "")}</span></td>')
        md_content.append(f"    </tr>")
        
        # 如果items是对象或有引用，递归展开（添加深度限制）
        if ((items_schema.get('type') == 'object' and 'properties' in items_schema) or 
            ('$ref' in items_schema and ref_name in definitions)) and indent_level < 4:  # 添加深度限制
            expand_schema_as_tree(md_content, items_schema, definitions, "", f"{parent_path}.{schema_name}", indent_level + 1, row_index)
        
        return
    
    # 处理对象类型
    if schema.get('type') == 'object' and 'properties' in schema:
        required_fields = schema.get('required', [])
        
        # 如果是顶层对象，不显示行，直接展开属性
        if indent_level == 0:
            for prop_name, prop_schema in schema['properties'].items():
                expand_schema_as_tree(md_content, prop_schema, definitions, prop_name, parent_path, indent_level, row_index)
        else:
            # 如果是嵌套对象，显示对象行并递归展开属性
            padding = indent_level * 20
            md_content.append(f"    <tr key={row_index}-{indent_level}>")
            md_content.append(f'      <td key=0><span style="padding-left: {padding}px"><span style="color: #8c8a8a">├─</span> {schema_name}</span></td>')
            md_content.append(f"      <td key=1><span>object</span></td>")
            md_content.append(f"      <td key=2>{'是' if schema_name in required_fields else '否'}</td>")
            md_content.append(f"      <td key=3></td>")
            md_content.append(f'      <td key=4><span style="white-space: pre-wrap">{schema.get("description", "")}</span></td>')
            md_content.append(f"    </tr>")
            
            # 递归展开属性（添加深度限制）
            if indent_level < 4:  # 添加深度限制
                for prop_name, prop_schema in schema['properties'].items():
                    expand_schema_as_tree(md_content, prop_schema, definitions, prop_name, f"{parent_path}.{schema_name}", indent_level + 1, row_index)
        
        return
    
    # 处理引用类型
    if '$ref' in schema:
        ref_name = schema['$ref'].split('/')[-1]
        padding = indent_level * 20
        
        # 显示引用行
        md_content.append(f"    <tr key={row_index}-{indent_level}>")
        md_content.append(f'      <td key=0><span style="padding-left: {padding}px"><span style="color: #8c8a8a">{"" if indent_level == 0 else "├─"}</span> {schema_name}</span></td>')
        md_content.append(f'      <td key=1><span>{ref_name}</span></td>')
        md_content.append(f"      <td key=2>{'是' if schema_name in (schema.get('required', [])) else '否'}</td>")
        md_content.append(f"      <td key=3>{schema.get('default', '')}</td>")
        md_content.append(f'      <td key=4><span style="white-space: pre-wrap">{schema.get("description", "")}</span></td>')
        md_content.append(f"    </tr>")
        
        # 如果引用存在于definitions中，递归展开
        if ref_name in definitions and indent_level < 3:  # 限制递归深度
            ref_schema = definitions[ref_name]
            expand_schema_as_tree(md_content, ref_schema, definitions, "", f"{parent_path}.{schema_name}", indent_level + 1, row_index)
        
        return
    
    # 处理基本类型
    prop_type = schema.get('type', '未知')
    if 'format' in schema:
        prop_type = f"{prop_type}({schema['format']})"
    
    padding = indent_level * 20
    
    # 显示基本类型行
    md_content.append(f"    <tr key={row_index}-{indent_level}>")
    md_content.append(f'      <td key=0><span style="padding-left: {padding}px"><span style="color: #8c8a8a">{"" if indent_level == 0 else "├─"}</span> {schema_name}</span></td>')
    md_content.append(f'      <td key=1><span>{prop_type}</span></td>')
    md_content.append(f"      <td key=2>{'是' if schema_name in (schema.get('required', [])) else '否'}</td>")
    md_content.append(f"      <td key=3>{schema.get('default', '')}</td>")
    md_content.append(f'      <td key=4><span style="white-space: pre-wrap">{schema.get("description", "")}</span></td>')
    md_content.append(f"    </tr>")


# ====== markdown_to_docx 工具类实现 ======
@register_tool('markdown_to_docx')
class MarkdownToDocxTool(BaseTool):
    """
    Markdown转Word工具
    """
    description = '将Markdown文件转换为Word文档'
    parameters = [{
        'name': 'md_file_path',
        'type': 'string',
        'description': 'Markdown文件的路径',
        'required': True
    }, {
        'name': 'output_file_path',
        'type': 'string',
        'description': '输出Word文档的路径（可选）',
        'required': False
    }]

    def call(self, params: str, **kwargs) -> str:
        try:
            args = json.loads(params)
            md_file_path = args['md_file_path']
            output_file_path = args.get('output_file_path')
            
            # 检查输入文件是否存在
            if not os.path.exists(md_file_path):
                return f"错误：Markdown文件 '{md_file_path}' 不存在"
            
            # 确定输出文件路径
            if not output_file_path:
                base_name = os.path.splitext(os.path.basename(md_file_path))[0]
                output_path = os.path.join(os.path.dirname(md_file_path or '.'), f"{base_name}.docx")
            else:
                output_path = output_file_path
            
            # 读取Markdown文件内容
            with open(md_file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 将Markdown转换为HTML
            html_content = markdown.markdown(md_content, extensions=[
                'extra',  # 额外功能，如表格、代码块等
                'toc',    # 目录生成
                'sane_lists',  # 更合理的列表处理
                'smarty'  # 智能引号等
            ])
            
            # 将HTML转换为Word文档
            # 创建Word文档
            doc = Document()
            
            # 解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 处理文档标题（如果有）
            title_tag = soup.find('h1')
            if title_tag:
                title_text = title_tag.get_text().strip()
                # 设置文档标题
                title_heading = doc.add_heading(title_text, level=0)
                title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 移除已处理的标题
                title_tag.decompose()
            
            # 遍历HTML内容的所有顶级元素
            for element in soup.find_all(recursive=False):
                # 处理标题
                if element.name.startswith('h') and len(element.name) == 2:
                    level = int(element.name[1])
                    text = element.get_text().strip()
                    
                    # 根据级别添加相应的标题
                    if 1 <= level <= 6:
                        # Word的heading级别是0-4，需要映射
                        word_level = min(4, level - 1)
                        doc.add_heading(text, level=word_level)
                
                # 处理段落
                elif element.name == 'p':
                    text = element.get_text().strip()
                    if text:
                        paragraph = doc.add_paragraph(text)
                        
                        # 检查段落是否有特殊样式
                        if element.get('align') == 'center':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif element.get('align') == 'right':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # 处理列表（有序和无序）
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        text = li.get_text().strip()
                        if text:
                            # 根据列表类型选择样式
                            style = 'List Bullet' if element.name == 'ul' else 'List Number'
                            doc.add_paragraph(text, style=style)
                
                # 处理表格
                elif element.name == 'table':
                    # 获取表头
                    thead = element.find('thead')
                    tbody = element.find('tbody')
                    
                    if thead and tbody:
                        # 获取表头行
                        header_row = thead.find('tr')
                        if header_row:
                            # 获取列数
                            header_cells = header_row.find_all(['th', 'td'])
                            col_count = len(header_cells)
                            
                            # 创建表格
                            table = doc.add_table(rows=1, cols=col_count)
                            
                            # 添加表头
                            for i, cell in enumerate(header_cells):
                                table.rows[0].cells[i].text = cell.get_text().strip()
                                # 设置表头样式（加粗）
                                for paragraph in table.rows[0].cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                            
                            # 添加表格数据
                            for row in tbody.find_all('tr'):
                                data_cells = row.find_all(['th', 'td'])
                                if data_cells:
                                    new_row = table.add_row()
                                    for i, cell in enumerate(data_cells[:col_count]):  # 防止列数不匹配
                                        new_row.cells[i].text = cell.get_text().strip()
                
                # 处理其他内容（简化实现）
                else:
                    text = element.get_text().strip()
                    if text:
                        doc.add_paragraph(text)
            
            # 保存Word文档
            doc.save(output_path)
            
            return f"Markdown文档已成功转换为Word文档：{output_path}"
        except Exception as e:
            return f"Markdown转Word失败: {str(e)}"

# ====== swagger_to_docx 工具类实现 ======
@register_tool('swagger_to_docx')
class SwaggerToDocxTool(BaseTool):
    """
    Swagger/OpenAPI JSON转Word工具（一步完成）
    """
    description = '将Swagger/OpenAPI JSON文件直接转换为Word文档'
    parameters = [{
        'name': 'swagger_file_path',
        'type': 'string',
        'description': 'Swagger/OpenAPI JSON文件的路径',
        'required': True
    }, {
        'name': 'output_file_path',
        'type': 'string',
        'description': '输出Word文档的路径（可选）',
        'required': False
    }]

    def call(self, params: str, **kwargs) -> str:
        try:
            args = json.loads(params)
            swagger_file_path = args['swagger_file_path']
            output_file_path = args.get('output_file_path')
            
            # 检查输入文件是否存在，如果不存在则尝试在GRADIO_TEMP_DIR目录下遍历查找
            if not os.path.exists(swagger_file_path):
                # 获取GRADIO_TEMP_DIR目录
                gradio_temp_dir = os.environ.get('GRADIO_TEMP_DIR', '')
                if gradio_temp_dir and os.path.exists(gradio_temp_dir):
                    # 获取要查找的文件名
                    target_filename = os.path.basename(swagger_file_path)
                    found_paths = []
                    
                    # 遍历GRADIO_TEMP_DIR目录及其所有子目录
                    for root, dirs, files in os.walk(gradio_temp_dir):
                        if target_filename in files:
                            found_paths.append(os.path.join(root, target_filename))
                    
                    # 根据查找结果处理
                    if len(found_paths) == 1:
                        swagger_file_path = found_paths[0]
                    elif len(found_paths) > 1:
                        return f"错误：在GRADIO_TEMP_DIR目录下找到多个同名文件：{', '.join(found_paths)}"
                    else:
                        return f"错误：Swagger文件 '{swagger_file_path}' 不存在，且在GRADIO_TEMP_DIR目录下也未找到"
                else:
                    return f"错误：Swagger文件 '{swagger_file_path}' 不存在"
            
            # 确定输出文件路径
            if not output_file_path:
                base_name = os.path.splitext(os.path.basename(swagger_file_path))[0]
                output_path = os.path.join(os.path.dirname(swagger_file_path or '.'), f"{base_name}_api文档.docx")
            else:
                output_path = output_file_path
            
            # 先生成临时Markdown文件
            temp_md_path = f"temp_{int(time.time()*1000)}.md"
            
            # 调用swagger_to_markdown工具
            swagger_tool = SwaggerToMarkdownTool()
            swagger_result = swagger_tool.call(json.dumps({
                'swagger_file_path': swagger_file_path,
                'output_file_path': temp_md_path
            }))
            
            if "成功" not in swagger_result:
                return f"Swagger转Word失败: {swagger_result}"
            
            # 调用markdown_to_docx工具
            md_tool = MarkdownToDocxTool()
            md_result = md_tool.call(json.dumps({
                'md_file_path': temp_md_path,
                'output_file_path': output_path
            }))
            
            # 删除临时文件
            if os.path.exists(temp_md_path):
                os.remove(temp_md_path)
            
            if "成功" in md_result:
                return f"Swagger文件已成功转换为Word文档：{output_path}"
            else:
                return f"Swagger转Word失败: {md_result}"
        except Exception as e:
            return f"Swagger转Word失败: {str(e)}"

# ====== 初始化API文档助手服务 ======
def init_agent_service():
    """初始化API文档助手服务"""
    llm_cfg = {
        'model': 'qwen-turbo-2025-04-28',
        'timeout': 30,
        'retry_count': 3,
    }
    try:
        bot = Assistant(
            llm=llm_cfg,
            name='项目API接口文档编写助手',
            description='将API接口编写为项目上的接口文档，采用Swagger/OpenAPI JSON转换为Word、Markdown文档',
            system_message=system_prompt,
            function_list=functions_desc,
        )
        print("助手初始化成功！")
        return bot
    except Exception as e:
        print(f"助手初始化失败: {str(e)}")
        raise

def app_gui():
    """图形界面模式，提供 Web 图形界面"""
    try:
        print("正在启动 Web 界面...")
        # 初始化助手
        bot = init_agent_service()
        # 配置聊天界面，列举3个典型API文档转换问题
        chatbot_config = {
            'prompt.suggestions': [
                '如何将Swagger JSON转换为Word文档？',
                '我有一个OpenAPI JSON文件，帮我转换为格式良好的API文档',
                '请将这个Swagger文件转换为Markdown格式',
            ]
        }
        print("Web 界面准备就绪，正在启动服务...")
        # 启动 Web 界面
        WebUI(
            bot,
            chatbot_config=chatbot_config
        ).run(server_name='0.0.0.0',server_port=7860)
    except Exception as e:
        print(f"启动 Web 界面失败: {str(e)}")
        print("请检查网络连接和 API Key 配置")


if __name__ == '__main__':
    # 运行模式选择
    app_gui()          # 图形界面模式（默认）